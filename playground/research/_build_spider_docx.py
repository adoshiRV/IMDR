"""Spider digest → branded .docx renderer (throwaway playground tooling).

Renders a Spider content MD (daily or weekly macro research digest) into a
self-contained RV-Capital-branded Word document. Generic markdown renderer —
Spider writes the MD, this turns it into the .docx — so there are no hard-coded
content blocks to keep in sync (unlike Jonah's renderer). The *styling* is
lifted from the JJ formatter: narrow margins, full-width green masthead banner,
serif section headings with rules, and thin-bordered tables with a green header
row, zebra striping and tight cell margins.

Supported markdown: h1/h2/h3 (`#`/`##`/`###`), bullet lists (`- `/`* `),
GFM pipe tables, `---` rules, paragraphs, inline `**bold**` / `*italic*` /
`code`, and a leading YAML front-matter block (used for the masthead subtitle).

Usage:
    python playground/research/_build_spider_docx.py IN.md [-o OUT.docx]

If -o is omitted the .docx is written next to the MD with the same stem.
This is playground demo tooling — not production, do not promote.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# --- RV brand palette (from rv_tokens.css / the JJ formatter) ---------------
GREEN = "004527"; DKBLUE = "001830"; LBLUE = "85A2BE"; LGREEN = "B2D0B9"
FG = "3D3E3E"; MUTED = "7A7C7C"; BORDER = "D9DCD8"; NEG = "B23A2B"
ZEBRA = "F1F5F2"; GREEN70 = "4D806A"; GREEN15 = "E0E9E4"
SERIF = "Georgia"; SANS = "Public Sans"

_REPO = Path(__file__).resolve().parents[2]
_ASSETS = _REPO / "docs/admin/research/brief_assets"
_LOGO_NEG = str(_ASSETS / "rv-logo-negative.png")


# ---- low-level OOXML helpers (ported from the JJ formatter) ----------------
def _shd(el, fill):
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), fill); el.append(s)


def cell_bg(cell, fill):
    _shd(cell._tc.get_or_add_tcPr(), fill)


def cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left),
                     ("end", right), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def no_borders(table):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); b.append(e)
    table._tbl.tblPr.append(b)


def thin_borders(table, color=BORDER, sz="4"):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz); e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color); b.append(e)
    table._tbl.tblPr.append(b)


def p_border_bottom(p, color=GREEN, sz="8"):
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz); bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color); pbdr.append(bottom); pPr.append(pbdr)


def keep_with_next(p):
    p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def repeat_header(row):
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(th)


def add_runs(p, text, *, size=9.5, color=FG, font=SANS, base_bold=False):
    """Add text honouring inline **bold**, *italic*, and `code` markers."""
    text = text.replace("`", "")
    for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if not tok:
            continue
        bold, ital, s = base_bold, False, tok
        if tok.startswith("**") and tok.endswith("**"):
            bold, s = True, tok[2:-2]
        elif tok.startswith("*") and tok.endswith("*"):
            ital, s = True, tok[1:-1]
        r = p.add_run(s); r.font.name = font; r.font.size = Pt(size)
        r.font.bold = bold; r.font.italic = ital
        r.font.color.rgb = RGBColor.from_string(color)


# ---- component renderers ---------------------------------------------------
def _front_matter(lines):
    """Split off a leading YAML block; return (meta dict, remaining lines)."""
    meta = {}
    if lines and lines[0].strip() == "---":
        end = next((i for i in range(1, len(lines)) if lines[i].strip() == "---"), None)
        if end is not None:
            for ln in lines[1:end]:
                if ":" in ln:
                    k, _, v = ln.partition(":")
                    meta[k.strip()] = v.strip()
            return meta, lines[end + 1:]
    return meta, lines


def masthead(doc, usable, title, subtitle):
    mast = doc.add_table(rows=1, cols=2); mast.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(mast)
    mast.columns[0].width = Inches(1.9); mast.columns[1].width = usable - Inches(1.9)
    lc, rc = mast.rows[0].cells
    lc.width = Inches(1.9); rc.width = usable - Inches(1.9)
    for c in (lc, rc):
        cell_bg(c, GREEN); cell_margins(c, 140, 140, 170, 170)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if Path(_LOGO_NEG).exists():
        lc.paragraphs[0].add_run().add_picture(_LOGO_NEG, width=Inches(1.55))
    tp = rc.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_after = Pt(0)
    r = tp.add_run(title); r.font.name = SERIF; r.font.size = Pt(19)
    r.font.bold = True; r.font.color.rgb = RGBColor.from_string("FFFFFF")
    if subtitle:
        sp = rc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sp.paragraph_format.space_before = Pt(2); sp.paragraph_format.space_after = Pt(0)
        r = sp.add_run(subtitle); r.font.name = SANS; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(LGREEN)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6); rule.paragraph_format.space_after = Pt(8)
    p_border_bottom(rule, GREEN, "12")


def heading(doc, text, *, serif_size=13.5):
    # split a trailing "*(...)*" tag into a muted run
    tag = ""
    m = re.search(r"\s*\*\((.+?)\)\*\s*$", text)
    if m:
        tag = m.group(1); text = text[:m.start()].strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    keep_with_next(p)
    rt = p.add_run(text.replace("`", "")); rt.font.name = SERIF
    rt.font.size = Pt(serif_size); rt.font.bold = True
    rt.font.color.rgb = RGBColor.from_string(GREEN)
    if tag:
        rr = p.add_run(f"   {tag.replace('`', '')}"); rr.font.name = SANS
        rr.font.size = Pt(8); rr.font.italic = True
        rr.font.color.rgb = RGBColor.from_string(MUTED)
    p_border_bottom(p, BORDER, "6")


def country_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    keep_with_next(p)
    r = p.add_run(text.replace("`", "")); r.font.name = SERIF
    r.font.size = Pt(11.5); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN70)
    p_border_bottom(p, GREEN15, "4")


def subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    keep_with_next(p)
    add_runs(p, text, size=9.5, color=DKBLUE, base_bold=True)


def bullet(doc, text, size=9.5):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(2.5); p.paragraph_format.line_spacing = 1.12
    sq = p.add_run("▪  "); sq.font.name = SANS; sq.font.size = Pt(8)
    sq.font.color.rgb = RGBColor.from_string(GREEN70)
    add_runs(p, text, size=size)


def paragraph(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    add_runs(p, text, size=9.5)


def _col_widths(rows, usable):
    ncol = max(len(r) for r in rows)
    maxlen = [1] * ncol
    for r in rows:
        for i in range(ncol):
            if i < len(r):
                maxlen[i] = max(maxlen[i], len(re.sub(r"[*`]", "", r[i])))
    clamped = [min(max(m, 6), 55) for m in maxlen]
    total = sum(clamped)
    return [Emu(int(usable * c / total)) for c in clamped]


def render_table(doc, usable, rows):
    if not rows:
        return
    headers, body = rows[0], rows[1:]
    ncol = max(len(r) for r in rows)
    cols = _col_widths(rows, usable)
    n_body = min(1 + len(body), 10)  # size shrinks a touch as tables widen
    body_size = 8.4 if ncol <= 6 else 7.8
    t = doc.add_table(rows=1, cols=ncol); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False; thin_borders(t)
    hdr = t.rows[0].cells; repeat_header(t.rows[0])
    for i in range(ncol):
        c = hdr[i]; c.width = cols[i]; cell_bg(c, GREEN); cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(1)
        pp.paragraph_format.space_after = Pt(1)
        txt = headers[i] if i < len(headers) else ""
        rr = pp.add_run(re.sub(r"[*`]", "", txt)); rr.font.name = SANS
        rr.font.size = Pt(body_size); rr.font.bold = True
        rr.font.color.rgb = RGBColor.from_string("FFFFFF")
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for ci in range(ncol):
            c = cells[ci]; c.width = cols[ci]; cell_margins(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 1:
                cell_bg(c, ZEBRA)
            pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(1.5)
            pp.paragraph_format.space_after = Pt(1.5)
            pp.paragraph_format.line_spacing = 1.08
            val = row[ci] if ci < len(row) else ""
            add_runs(pp, str(val), size=body_size, base_bold=(ci == 0))


def footer_band(doc, sec, usable, title_line, sub_line):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    foot = doc.add_table(rows=1, cols=2); foot.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(foot)
    foot.columns[0].width = Inches(1.7); foot.columns[1].width = usable - Inches(1.7)
    fl, fr = foot.rows[0].cells
    fl.width = Inches(1.7); fr.width = usable - Inches(1.7)
    for c in (fl, fr):
        cell_bg(c, DKBLUE); cell_margins(c, 150, 150, 170, 170)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if Path(_LOGO_NEG).exists():
        fl.paragraphs[0].add_run().add_picture(_LOGO_NEG, width=Inches(1.4))
    ap = fr.paragraphs[0]; ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = ap.add_run(title_line); r.font.name = SANS; r.font.size = Pt(9)
    r.font.bold = True; r.font.color.rgb = RGBColor.from_string("FFFFFF")
    ap2 = fr.add_paragraph(); ap2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ap2.paragraph_format.space_before = Pt(2)
    r = ap2.add_run(sub_line); r.font.name = SANS; r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(LBLUE)
    # running footer with page number
    fpara = sec.footer.paragraphs[0]; fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = fpara.add_run(title_line + "  ·  page ")
    rr.font.name = SANS; rr.font.size = Pt(7.5)
    rr.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "15"); rpr.append(sz)
    run_el.append(rpr); t_el = OxmlElement("w:t"); t_el.text = "1"
    run_el.append(t_el); fld.append(run_el); fpara._p.append(fld)


# ---- main render -----------------------------------------------------------
def render(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    meta, lines = _front_matter(lines)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.55); sec.right_margin = Inches(0.55)
    usable = sec.page_width - sec.left_margin - sec.right_margin
    stl = doc.styles["Normal"]
    stl.font.name = SANS; stl.font.size = Pt(9.5)
    stl.font.color.rgb = RGBColor.from_string(FG)
    stl.paragraph_format.space_after = Pt(4); stl.paragraph_format.line_spacing = 1.18

    title = next((ln[2:].strip() for ln in lines if ln.startswith("# ")),
                 "Macro Research Digest")
    edition = meta.get("edition", "").capitalize()
    date = meta.get("date", "")
    bits = [b for b in ("Rates / FX desk", date, (edition + " · DRAFT" if edition else "")) if b]
    masthead(doc, usable, re.sub(r"[*`]", "", title), "  ·  ".join(bits))

    tbl_buf: list[list[str]] = []
    first_h1_seen = False

    def flush():
        nonlocal tbl_buf
        render_table(doc, usable, tbl_buf)
        tbl_buf = []

    for ln in lines:
        raw = ln.rstrip()
        stripped = raw.strip()
        # GFM table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if cells and all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            tbl_buf.append(cells)
            continue
        flush()

        if not stripped:
            continue
        if raw.startswith("### "):
            country_heading(doc, raw[4:].strip())
        elif raw.startswith("## "):
            heading(doc, raw[3:].strip())
        elif raw.startswith("# "):
            if first_h1_seen:
                heading(doc, raw[2:].strip(), serif_size=15)
            first_h1_seen = True
        elif stripped in ("---", "***", "___"):
            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(4)
            rule.paragraph_format.space_after = Pt(4)
            p_border_bottom(rule, BORDER, "6")
        elif re.match(r"^\s*[-*]\s+", raw):
            bullet(doc, re.sub(r"^\s*[-*]\s+", "", raw))
        elif re.match(r"^\*\*[A-Z]\.", stripped) or (stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 90):
            subhead(doc, stripped)  # e.g. "**A. Themes in play**"
        else:
            paragraph(doc, stripped)

    flush()
    ed = edition or "Macro"
    footer_band(
        doc, sec, usable,
        f"RV Capital  ·  {ed} Macro Research Digest (Spider · demo)",
        "Opinion & trade-extraction tool, not a clipping service. "
        "Internal use only — not investment advice.",
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {out_path}  ({out_path.stat().st_size} bytes)")


def main() -> None:
    ap = argparse.ArgumentParser(description="Render a Spider digest MD to branded .docx")
    ap.add_argument("md", type=Path, help="input content markdown")
    ap.add_argument("-o", "--out", type=Path, default=None, help="output .docx path")
    args = ap.parse_args()
    out = args.out or args.md.with_suffix(".docx")
    render(args.md, out)


if __name__ == "__main__":
    main()
